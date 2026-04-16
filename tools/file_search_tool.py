from __future__ import annotations

import hashlib
import json
import os
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Optional

import pandas as pd

from config import DATA_DIR

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

CACHE_FILE = DATA_DIR / "cache_busca_v3.json"
EXTENSOES_SUPORTADAS = {".txt", ".docx", ".xlsx", ".xlsm", ".pdf"}


@dataclass
class SearchResult:
    arquivo: str
    tipo: str
    local: str
    conteudo: str
    status: str = "OK"


@dataclass
class SearchConfig:
    termo: str
    pasta: str
    saida: str
    max_workers: int = min(32, (os.cpu_count() or 4) * 2)


class SearchCancellationToken:
    def __init__(self) -> None:
        self._event = threading.Event()

    def cancel(self) -> None:
        self._event.set()

    @property
    def is_cancelled(self) -> bool:
        return self._event.is_set()


def build_output_path(base_dir: Path | None = None) -> str:
    destino = base_dir or DATA_DIR
    nome = f"resultado_busca_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return str(destino / nome)


def gerar_assinatura_pasta(pasta: str) -> str:
    h = hashlib.md5()
    for path in sorted(Path(pasta).rglob("*")):
        if not path.is_file() or path.suffix.lower() not in EXTENSOES_SUPORTADAS:
            continue
        try:
            stat = path.stat()
            payload = f"{path}|{stat.st_size}|{int(stat.st_mtime)}"
            h.update(payload.encode("utf-8", errors="ignore"))
        except OSError:
            h.update(str(path).encode("utf-8", errors="ignore"))
    return h.hexdigest()


def extrair_snippet(texto: str, termo: str, janela: int = 80) -> str:
    texto_lower = texto.lower()
    termo_lower = termo.lower()
    idx = texto_lower.find(termo_lower)
    if idx == -1:
        return texto.strip()[:200]
    ini = max(0, idx - janela)
    fim = min(len(texto), idx + len(termo) + janela)
    trecho = texto[ini:fim].replace("\n", " ").strip()
    return f"...{trecho}..." if ini > 0 or fim < len(texto) else trecho


def buscar_txt(file: Path, termo: str) -> list[SearchResult]:
    resultados: list[SearchResult] = []
    try:
        with file.open("r", encoding="utf-8", errors="ignore") as f:
            for i, linha in enumerate(f, start=1):
                if termo.lower() in linha.lower():
                    resultados.append(SearchResult(str(file), "txt", f"Linha {i}", linha.strip()))
    except Exception as exc:
        resultados.append(SearchResult(str(file), "txt", "Erro", str(exc), "ERRO"))
    return resultados


def buscar_docx(file: Path, termo: str) -> list[SearchResult]:
    if Document is None:
        return [SearchResult(str(file), "docx", "Erro", "python-docx não instalado", "ERRO")]
    resultados: list[SearchResult] = []
    try:
        doc = Document(file)
        for i, p in enumerate(doc.paragraphs, start=1):
            texto = (p.text or "").strip()
            if texto and termo.lower() in texto.lower():
                resultados.append(SearchResult(str(file), "docx", f"Parágrafo {i}", texto))
        for t_idx, table in enumerate(doc.tables, start=1):
            for r_idx, row in enumerate(table.rows, start=1):
                for c_idx, cell in enumerate(row.cells, start=1):
                    texto = (cell.text or "").strip()
                    if texto and termo.lower() in texto.lower():
                        resultados.append(SearchResult(str(file), "docx", f"Tabela {t_idx} - L{r_idx}C{c_idx}", texto))
    except Exception as exc:
        resultados.append(SearchResult(str(file), "docx", "Erro", str(exc), "ERRO"))
    return resultados


def buscar_excel(file: Path, termo: str) -> list[SearchResult]:
    if load_workbook is None:
        return [SearchResult(str(file), "excel", "Erro", "openpyxl não instalado", "ERRO")]
    resultados: list[SearchResult] = []
    try:
        wb = load_workbook(filename=file, read_only=True, data_only=True)
        termo_lower = termo.lower()
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    texto = str(cell.value)
                    if termo_lower in texto.lower():
                        resultados.append(SearchResult(str(file), "excel", f"{ws.title}!{cell.coordinate}", texto))
        wb.close()
    except Exception as exc:
        resultados.append(SearchResult(str(file), "excel", "Erro", str(exc), "ERRO"))
    return resultados


def buscar_pdf(file: Path, termo: str) -> list[SearchResult]:
    if fitz is None:
        return [SearchResult(str(file), "pdf", "Erro", "PyMuPDF não instalado", "ERRO")]
    resultados: list[SearchResult] = []
    try:
        with fitz.open(file) as pdf:
            for i in range(pdf.page_count):
                page = pdf.load_page(i)
                texto = page.get_text("text") or ""
                if termo.lower() in texto.lower():
                    resultados.append(SearchResult(str(file), "pdf", f"Página {i + 1}", extrair_snippet(texto, termo)))
    except Exception as exc:
        resultados.append(SearchResult(str(file), "pdf", "Erro", str(exc), "ERRO"))
    return resultados

FUNCOES_BUSCA: dict[str, Callable[[Path, str], list[SearchResult]]] = {
    ".txt": buscar_txt,
    ".docx": buscar_docx,
    ".xlsx": buscar_excel,
    ".xlsm": buscar_excel,
    ".pdf": buscar_pdf,
}


def processar_arquivo(file: Path, termo: str) -> list[SearchResult]:
    func = FUNCOES_BUSCA.get(file.suffix.lower())
    return func(file, termo) if func else []


def salvar_cache(chave: str, termo: str, resultados: list[SearchResult]) -> None:
    try:
        data = {"hash": chave, "termo": termo, "resultados": [asdict(r) for r in resultados]}
        CACHE_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


def carregar_cache(chave: str, termo: str) -> Optional[list[SearchResult]]:
    try:
        if not CACHE_FILE.exists():
            return None
        data = json.loads(CACHE_FILE.read_text(encoding="utf-8"))
        if data.get("hash") == chave and data.get("termo") == termo:
            return [SearchResult(**item) for item in data.get("resultados", [])]
    except Exception:
        return None
    return None


def summarize_results(config: SearchConfig, resultados: list[SearchResult], origem_cache: bool, cancelled: bool = False) -> str:
    total = len(resultados)
    erros = sum(1 for item in resultados if item.status == "ERRO")
    arquivos_unicos = len({item.arquivo for item in resultados})
    tipos: dict[str, int] = {}
    for item in resultados:
        tipos[item.tipo] = tipos.get(item.tipo, 0) + 1
    tipos_texto = ", ".join(f"{tipo}:{quantidade}" for tipo, quantidade in sorted(tipos.items())) or "sem ocorrências"
    linhas = [
        "Busca cancelada." if cancelled else "Busca concluída.",
        f"Termo: {config.termo}",
        f"Pasta: {config.pasta}",
        f"Saída: {config.saida}",
        f"Registros: {total} | Arquivos com ocorrência: {arquivos_unicos} | Erros: {erros}",
        f"Origem: {'Cache' if origem_cache else 'Busca nova'}",
        f"Tipos encontrados: {tipos_texto}",
    ]
    exemplos = [item for item in resultados if item.status == "OK"][:5]
    if exemplos:
        linhas.append("Amostra:")
        for item in exemplos:
            linhas.append(f"- [{item.tipo}] {item.local} | {Path(item.arquivo).name}")
    return "\n".join(linhas)


class FileSearchTool:
    def search(
        self,
        config: SearchConfig,
        progress_callback: Optional[Callable[[int, int, Path, int, int], None]] = None,
        log_callback: Optional[Callable[[str], None]] = None,
        cancellation_token: SearchCancellationToken | None = None,
    ) -> tuple[list[SearchResult], bool, bool]:
        arquivos = [
            f
            for f in Path(config.pasta).rglob("*")
            if f.is_file() and f.suffix.lower() in EXTENSOES_SUPORTADAS
        ]
        if not arquivos:
            return [], False, False

        assinatura = gerar_assinatura_pasta(config.pasta)
        cache = carregar_cache(assinatura, config.termo)
        if cache is not None:
            if log_callback:
                log_callback("Cache válido encontrado. Reutilizando resultados.")
            if progress_callback:
                progress_callback(len(arquivos), len(arquivos), Path(config.pasta), len(cache), 0)
            return cache, True, False

        resultados: list[SearchResult] = []
        erros = 0
        total = len(arquivos)
        cancelled = False

        with ThreadPoolExecutor(max_workers=config.max_workers) as executor:
            futures = {executor.submit(processar_arquivo, arq, config.termo): arq for arq in arquivos}
            for idx, future in enumerate(as_completed(futures), start=1):
                if cancellation_token and cancellation_token.is_cancelled:
                    cancelled = True
                    if log_callback:
                        log_callback("Cancelamento solicitado. Encerrando processamento restante...")
                    for pending in futures:
                        pending.cancel()
                    executor.shutdown(wait=False, cancel_futures=True)
                    break
                arquivo = futures[future]
                try:
                    achados = future.result()
                    resultados.extend(achados)
                    if any(r.status == "ERRO" for r in achados):
                        erros += 1
                        if log_callback:
                            log_callback(f"[ERRO] {arquivo.name}")
                    elif log_callback:
                        log_callback(f"[OK] {arquivo.name}")
                except Exception as exc:
                    erros += 1
                    resultados.append(SearchResult(str(arquivo), arquivo.suffix.lower(), "Erro", str(exc), "ERRO"))
                    if log_callback:
                        log_callback(f"[ERRO CRÍTICO] {arquivo.name}: {exc}")
                if progress_callback:
                    progress_callback(idx, total, arquivo, len(resultados), erros)

        if not cancelled:
            salvar_cache(assinatura, config.termo, resultados)
        return resultados, False, cancelled

    def export_results(
        self,
        config: SearchConfig,
        resultados: list[SearchResult],
        origem_cache: bool,
        cancelled: bool = False,
    ) -> str:
        df = pd.DataFrame(
            [asdict(r) for r in resultados],
            columns=["arquivo", "tipo", "local", "conteudo", "status"],
        )
        with pd.ExcelWriter(config.saida, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="Resultados", index=False)
            resumo = pd.DataFrame(
                {
                    "Métrica": ["Termo", "Pasta", "Total de registros", "Erros", "Origem", "Cancelado"],
                    "Valor": [
                        config.termo,
                        config.pasta,
                        len(df),
                        int((df["status"] == "ERRO").sum()) if not df.empty else 0,
                        "Cache" if origem_cache else "Busca nova",
                        "Sim" if cancelled else "Não",
                    ],
                }
            )
            resumo.to_excel(writer, sheet_name="Resumo", index=False)
        return config.saida
